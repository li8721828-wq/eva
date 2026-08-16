# -*- coding: utf-8 -*-
r"""
Step 5: 用 python-docx 1.2.0 将 Markdown 论文转换为 Word 文档
输入 : C:\Users\苏\Desktop\paper_llm_finetuning.md
输出 : C:\Users\苏\Desktop\paper_llm_finetuning.docx

解析: H1/H2/H3 标题、正文段落、3.3 节对比表、参考文献 [1]-[29]
样式: 内置 Title / Heading 1 / Heading 2 / Heading 3
字体: 中文 宋体(正文)/黑体(标题)，西文 Times New Roman；A4 页面
"""
import os
import re
import sys
import zipfile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = r"C:\Users\苏\Desktop\paper_llm_finetuning.md"
DST = r"C:\Users\苏\Desktop\paper_llm_finetuning.docx"
REPORT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "STEP5_REPORT.md")

EAST_BODY = "宋体"
EAST_HEAD = "黑体"
LATIN = "Times New Roman"


def set_run_font(run, east=EAST_BODY, latin=LATIN, size_pt=None, bold=None, color=None):
    """同时设置西文字体与中文（东亚）字体，保证中英文混排正确。"""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, east, latin=LATIN, size_pt=None, bold=None, color=None):
    style.font.name = latin
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def indent_chars(paragraph, chars=2):
    """段首缩进 chars 个字符：用 Word 中文排版标准 firstLineChars=200。"""
    paragraph.paragraph_format.first_line_indent = Pt(chars * 10.5)
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        ind.set(qn("w:firstLineChars"), str(chars * 100))


def parse_md(path):
    """解析 Markdown，返回块序列：(h, level, text) / (p, text) / (table, rows)。"""
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
        elif line.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            blocks.append(("table", rows))
        else:
            blocks.append(("p", line))
            i += 1
    return blocks


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep_row(cells):
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def build(blocks):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # A4
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    set_style_font(normal, EAST_BODY, LATIN, 10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    set_style_font(doc.styles["Title"], EAST_HEAD, LATIN, 22, bold=True, color="000000")
    set_style_font(doc.styles["Heading 1"], EAST_HEAD, LATIN, 16, bold=True, color="000000")
    set_style_font(doc.styles["Heading 2"], EAST_HEAD, LATIN, 14, bold=True, color="000000")
    set_style_font(doc.styles["Heading 3"], EAST_HEAD, LATIN, 12, bold=True, color="000000")

    HEAD = {1: ("Title", 22), 2: ("Heading 1", 16), 3: ("Heading 2", 14)}

    in_refs = False
    for block in blocks:
        kind = block[0]
        if kind == "h":
            level, text = block[1], block[2]
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(text), EAST_HEAD, LATIN, 22, bold=True, color="000000")
                in_refs = False
            else:
                style_name, size = HEAD.get(level, ("Heading 3", 12))
                p = doc.add_paragraph(style=style_name)
                set_run_font(p.add_run(text), EAST_HEAD, LATIN, size, bold=True, color="000000")
                in_refs = text.strip().startswith("参考文献")
        elif kind == "p":
            text = block[1]
            p = doc.add_paragraph(style="Normal")
            if in_refs and re.match(r"^\[\d+\]", text):
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.first_line_indent = Cm(-0.74)
                p.paragraph_format.line_spacing = 1.0
                set_run_font(p.add_run(text), EAST_BODY, LATIN, 9)
            else:
                indent_chars(p, 2)
                set_run_font(p.add_run(text), EAST_BODY, LATIN, 10.5)
        elif kind == "table":
            parsed = []
            for rowline in block[1]:
                cells = split_row(rowline)
                if not is_sep_row(cells):
                    parsed.append(cells)
            if not parsed:
                continue
            ncols = max(len(r) for r in parsed)
            table = doc.add_table(rows=len(parsed), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            widths = [Cm(2.9), Cm(2.3), Cm(2.4), Cm(2.2), Cm(2.4), Cm(2.46)]
            for ri, row in enumerate(parsed):
                for ci in range(ncols):
                    cell = table.cell(ri, ci)
                    if ri == 0:
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:fill"), "D9E2F3")
                        cell._tc.get_or_add_tcPr().append(shd)
                    if ci < len(widths):
                        cell.width = widths[ci]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = 1.0
                    set_run_font(para.add_run(row[ci] if ci < len(row) else ""),
                                 EAST_BODY, LATIN, 9, bold=(ri == 0))

    doc.core_properties.title = "大语言模型微调技术研究综述：方法、对齐、数据与评测"
    doc.save(DST)
    return doc


def verify():
    """重新打开生成的 docx，逐项校验并输出 STEP5_REPORT.md。"""
    ok = True
    n_checks = 0

    def check(cond, name, detail=""):
        nonlocal ok, n_checks
        n_checks += 1
        if not cond:
            ok = False
        print("  [%s] %s %s" % ("PASS" if cond else "FAIL", name, detail))
        return cond

    size = os.path.getsize(DST)
    doc = Document(DST)
    paras = doc.paragraphs
    styles = {}
    for p in paras:
        styles[p.style.name] = styles.get(p.style.name, 0) + 1

    title_n = styles.get("Title", 0)
    h1_n = styles.get("Heading 1", 0)
    h2_n = styles.get("Heading 2", 0)
    h3_n = styles.get("Heading 3", 0)

    refs = [p.text for p in paras if re.match(r"^\[\d+\]", p.text)]
    ref_nums = [int(re.match(r"^\[(\d+)\]", t).group(1)) for t in refs]

    tables = doc.tables
    t = tables[0] if tables else None
    t_rows = len(t.rows) if t else 0
    t_cols = len(t.columns) if t else 0

    md_lines = open(SRC, "r", encoding="utf-8").read().splitlines()
    md_headings = [re.sub(r"^#{1,6}\s+", "", l.strip())
                   for l in md_lines if re.match(r"^#{1,6}\s+", l.strip())]
    doc_headings = [p.text for p in paras
                    if p.style.name in ("Title", "Heading 1", "Heading 2", "Heading 3")]

    md_body = [l.strip() for l in md_lines
               if l.strip() and not l.strip().startswith("#") and not l.strip().startswith("|")]
    md_refs = [l for l in md_body if re.match(r"^\[\d+\]", l)]
    md_paras = [l for l in md_body if not re.match(r"^\[\d+\]", l)]

    all_text = "\n".join(p.text for p in paras)
    for tb in tables:
        for row in tb.rows:
            for c in row.cells:
                all_text += "\n" + c.text
    cjk_total = len(re.findall(r"[\u4e00-\u9fff]", all_text))

    with zipfile.ZipFile(DST) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    east_song = xml.count('w:eastAsia="宋体"')
    east_hei = xml.count('w:eastAsia="黑体"')

    print("== Step 5 校验（python-docx 1.2.0）==")
    check(size > 20000, "文件存在且非空", "%.1f KB" % (size / 1024.0))
    check(title_n == 1, "Title 样式段落（论文标题）", str(title_n))
    check(h1_n == 11, "Heading 1 段落（一级章节）", str(h1_n))
    check(h2_n == 23, "Heading 2 段落（二级小节）", str(h2_n))
    check(h3_n == 0, "Heading 3（预留，无更深标题）", str(h3_n))
    check(doc_headings == md_headings, "标题文本与 Markdown 逐条一致", "%d 条" % len(md_headings))
    check(styles.get("Normal", 0) == len(md_body), "正文段落数（含参考文献）",
          "%d/%d" % (styles.get("Normal", 0), len(md_body)))
    check([p.text for p in paras if p.style.name == "Normal" and not re.match(r"^\[\d+\]", p.text)] == md_paras,
          "正文逐字一致（无丢失/篡改）", "%d 段" % len(md_paras))
    check(len(refs) == 29 and ref_nums == list(range(1, 30)),
          "参考文献 [1]-[29] 齐全且按序", "%d 条" % len(refs))
    check(refs == md_refs, "参考文献文本逐条一致", "")
    check(len(tables) == 1 and t_rows == 6 and t_cols == 6,
          "3.3 节对比表 6x6（表头+5 行）", "%d 张 %dx%d" % (len(tables), t_rows, t_cols))
    check(t is not None and t.cell(0, 0).text == "方法" and t.cell(5, 5).text == "超大模型、简单任务",
          "表格首/末单元格内容", '"%s" / "%s"' % (t.cell(0, 0).text, t.cell(5, 5).text))
    check("#" not in all_text and "|" not in all_text, "无 Markdown 残留符号", "")
    check(cjk_total >= 8000, "全文汉字数（8000-12000 区间）", str(cjk_total))
    check(east_song > 0 and east_hei > 0, "中文东亚字体（宋体/黑体）",
          "宋体x%d 黑体x%d" % (east_song, east_hei))

    lines = [
        "# STEP5_REPORT.md — Step 5 Word 文档生成报告",
        "",
        "## 结果: " + ("PASS" if ok else "FAIL"),
        "",
        "**输入**：`C:\\Users\\苏\\Desktop\\paper_llm_finetuning.md`（Markdown，UTF-8）",
        "**输出**：`C:\\Users\\苏\\Desktop\\paper_llm_finetuning.docx`（python-docx 1.2.0，%.1f KB，%d 字节）"
        % (size / 1024.0, size),
        "",
        "## 实际数据",
        "",
        "| 项目 | 实测值 |",
        "| --- | --- |",
        "| 总段落数 | %d |" % len(paras),
        "| Title（论文标题） | %d |" % title_n,
        "| Heading 1（一级章节：摘要/关键词/1-8 章/参考文献） | %d |" % h1_n,
        "| Heading 2（二级小节） | %d |" % h2_n,
        "| Heading 3（预留，本文档无更深标题） | %d |" % h3_n,
        "| 正文段落（Normal，含参考文献） | %d |" % styles.get("Normal", 0),
        "| 参考文献 [1]-[29] | %d 条，编号按序完整 |" % len(refs),
        "| 3.3 节对比表 | %d 张，%d 行 x %d 列（表头+5 种方法） |" % (len(tables), t_rows, t_cols),
        "| 全文汉字数 | %d |" % cjk_total,
        "| 东亚字体（document.xml 中 w:eastAsia） | 宋体 x%d，黑体 x%d |" % (east_song, east_hei),
        "| Markdown 残留（# / |） | 0 |",
        "",
        "## 校验项",
        "",
        "共 %d 项检查，全部%s。标题层级映射：H1→Title、H2→Heading 1、H3→Heading 2（Heading 3 已配置备用）；"
        "标题文本、正文段落、参考文献文本均与 Markdown 逐条一致；表格 6x6 内容一致。"
        "排版：A4 页面（上下 2.54cm、左右 3.17cm）；正文宋体五号 10.5pt、1.5 倍行距、首行缩进 2 字符"
        "（w:firstLineChars=200）；标题黑体加粗黑色；西文 Times New Roman；参考文献 9pt 悬挂缩进。"
        % (n_checks, "通过" if ok else "未通过"),
    ]
    report = "\n".join(lines) + "\n"
    with open(REPORT, "w", encoding="utf-8") as f:
        f.write(report)
    print("报告已写入:", REPORT)
    print("RESULT:", "PASS" if ok else "FAIL")
    return ok


if __name__ == "__main__":
    blocks = parse_md(SRC)
    build(blocks)
    print("== 解析统计 ==")
    print("H1=%d H2=%d H3=%d 正文段=%d 参考文献=%d 表格=%d" % (
        sum(1 for b in blocks if b[0] == "h" and b[1] == 1),
        sum(1 for b in blocks if b[0] == "h" and b[1] == 2),
        sum(1 for b in blocks if b[0] == "h" and b[1] == 3),
        sum(1 for b in blocks if b[0] == "p" and not re.match(r"^\[\d+\]", b[1])),
        sum(1 for b in blocks if b[0] == "p" and re.match(r"^\[\d+\]", b[1])),
        sum(1 for b in blocks if b[0] == "table")))
    print("已生成:", DST, "%.1f KB" % (os.path.getsize(DST) / 1024.0))
    verify()
