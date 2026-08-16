# -*- coding: utf-8 -*-
"""Step 5 独立回读校验：重新打开生成的 docx，输出关键实测数字（不复用生成脚本的校验逻辑）。"""
import re
import sys

from docx import Document

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

DST = r"C:\Users\苏\Desktop\paper_llm_finetuning.docx"

doc = Document(DST)
styles = {}
for p in doc.paragraphs:
    styles[p.style.name] = styles.get(p.style.name, 0) + 1

refs = [p.text for p in doc.paragraphs if re.match(r"^\[\d+\]", p.text)]
ref_nums = [int(re.match(r"^\[(\d+)\]", t).group(1)) for t in refs]
t = doc.tables[0]

text = "\n".join(p.text for p in doc.paragraphs)
for tb in doc.tables:
    for row in tb.rows:
        for c in row.cells:
            text += "\n" + c.text
cjk = len(re.findall(r"[\u4e00-\u9fff]", text))

print("== 独立回读（step5_readback.py）==")
print("paragraphs total:", len(doc.paragraphs))
print("styles:", dict(sorted(styles.items())))
print("headings Title/H1/H2:",
      sum(1 for p in doc.paragraphs if p.style.name == "Title"),
      sum(1 for p in doc.paragraphs if p.style.name == "Heading 1"),
      sum(1 for p in doc.paragraphs if p.style.name == "Heading 2"))
print("refs:", len(refs), "| in order:", ref_nums == list(range(1, 30)))
print("first ref:", refs[0][:36])
print("last ref:", refs[-1][:36])
print("table:", len(t.rows), "rows x", len(t.columns), "cols")
print("table header:", [c.text for c in t.rows[0].cells])
print("table last row:", [c.text for c in t.rows[5].cells])
print("first para:", repr(doc.paragraphs[0].text[:44]), "| style:", doc.paragraphs[0].style.name)
print("last para:", repr(doc.paragraphs[-1].text[:60]))
print("CJK total:", cjk)
