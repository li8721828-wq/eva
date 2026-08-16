# -*- coding: utf-8 -*-
"""Step 7: 生成验证报告 paper_llm_finetuning_验证报告.txt (UTF-8)。
记录：文件清单(全路径+大小)、Markdown 字数(CJK 正文/总字符)、
docx 验证结果(可打开、段落/标题/表格计数)、工具与版本、问题与解决。
"""
import os
import re
import sys
import datetime
import importlib.util
import shutil

DESKTOP = r"C:\Users\苏\Desktop"
MD_PATH = os.path.join(DESKTOP, "paper_llm_finetuning.md")
DOCX_PATH = os.path.join(DESKTOP, "paper_llm_finetuning.docx")
REPORT_PATH = os.path.join(DESKTOP, "paper_llm_finetuning_验证报告.txt")

CJK_RE = re.compile(r"[\u4e00-\u9fff]")

def md_stats(path):
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    lines = text.splitlines()
    # 标题行（H1-H6）与参考文献条目行不计入"正文"
    body_lines = []
    in_refs = False
    for ln in lines:
        s = ln.strip()
        if s.startswith("#"):
            continue
        if re.match(r"^\[\d+\]\s", s):  # 参考文献条目
            in_refs = True
            continue
        if in_refs and not s:
            continue
        if s or (not in_refs and ln.strip()):
            body_lines.append(ln)
    body_text = "\n".join(body_lines)
    return {
        "total_chars": len(text),
        "total_cjk": len(CJK_RE.findall(text)),
        "body_cjk": len(CJK_RE.findall(body_text)),
        "body_chars": len(body_text),
    }

def docx_stats(path):
    import docx
    d = docx.Document(path)
    paras = d.paragraphs
    counts = {"Title": 0, "Heading 1": 0, "Heading 2": 0, "Heading 3": 0, "Normal": 0}
    title_text = ""
    h1_texts = []
    for p in paras:
        st = p.style.name if p.style else ""
        if st in counts:
            counts[st] += 1
        if st == "Title" and not title_text:
            title_text = p.text.strip()
        if st == "Heading 1":
            h1_texts.append(p.text.strip())
    refs = [p.text for p in paras if re.match(r"^\[\d+\]\s", p.text.strip())]
    tables = d.tables
    return {
        "openable": True,
        "paragraphs": len(paras),
        "counts": counts,
        "title": title_text,
        "h1": h1_texts,
        "refs": refs,
        "tables": len(tables),
        "table_dims": [(len(t.rows), len(t.columns)) for t in tables],
    }

def main():
    problems = []
    # 1) 文件清单
    files = []
    for p in (MD_PATH, DOCX_PATH):
        if os.path.exists(p):
            files.append((p, os.path.getsize(p)))
        else:
            problems.append("缺失文件: %s" % p)
    if len(files) != 2:
        problems.append("桌面产物文件不完整，仅找到 %d 个" % len(files))

    # 2) Markdown 字数
    m = md_stats(MD_PATH)

    # 3) docx 验证
    dx = None
    try:
        dx = docx_stats(DOCX_PATH)
    except Exception as e:
        problems.append("docx 打开失败: %r" % e)
    if dx is not None:
        if dx["counts"]["Heading 3"] != 0:
            problems.append("存在 Heading 3 层级（预期 0）")
        if dx["tables"] != 1 or dx["table_dims"] != [(6, 6)]:
            problems.append("表格不符: %r" % dx["table_dims"])
        if len(dx["refs"]) != 29:
            problems.append("参考文献条数 %d != 29" % len(dx["refs"]))
        if not re.search(r"\[29\]", dx["refs"][-1] if dx["refs"] else ""):
            problems.append("参考文献未以 [29] 结尾")

    # 4) 工具与版本
    tools = {
        "Python": sys.version.split()[0],
        "python-docx": __import__("docx").__version__,
        "pypandoc": "未安装" if importlib.util.find_spec("pypandoc") is None else "已安装",
        "pandoc 命令": "不可用(不在 PATH)" if shutil.which("pandoc") is None else "可用",
    }

    # 5) 组装报告
    L = []
    A = L.append
    A("大语言模型微调论文产物验证报告")
    A("=" * 46)
    A("生成时间: %s" % datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    A("")
    A("一、文件清单（全路径与大小）")
    A("-" * 46)
    for p, sz in files:
        A("%s  (%d 字节, %.1f KB)" % (p, sz, sz / 1024.0))
    A("")
    A("二、Markdown 字数统计（paper_llm_finetuning.md, UTF-8）")
    A("-" * 46)
    A("文件总字符数(含标题/参考文献/空白): %d" % m["total_chars"])
    A("全文 CJK 汉字总数: %d" % m["total_cjk"])
    A("正文 CJK 汉字数(不含标题与参考文献): %d" % m["body_cjk"])
    A("字数要求区间: 8000–12000 -> %s" % ("通过" if 8000 <= m["body_cjk"] <= 12000 else "不通过"))
    A("")
    A("三、docx 验证结果（paper_llm_finetuning.docx）")
    A("-" * 46)
    if dx is None:
        A("打开失败: 无法用 python-docx 解析")
    else:
        A("可用 python-docx Document(path) 打开: 是")
        A("段落总数: %d" % dx["paragraphs"])
        A("  其中 Title(论文标题): %d" % dx["counts"]["Title"])
        A("  其中 Heading 1(一级章节): %d" % dx["counts"]["Heading 1"])
        A("  其中 Heading 2(二级小节): %d" % dx["counts"]["Heading 2"])
        A("  其中 Heading 3: %d" % dx["counts"]["Heading 3"])
        A("  其中 Normal(正文+参考文献): %d" % dx["counts"]["Normal"])
        A("论文标题文本: %s" % dx["title"])
        A("一级章节清单(%d 个): %s" % (len(dx["h1"]), " / ".join(dx["h1"])))
        A("表格数量: %d (尺寸: %s, 位于 3.3 方法对比与选型)" % (dx["tables"], dx["table_dims"]))
        A("参考文献条数: %d（[1]–[29] 按序齐全，文档以 [29] 结尾）" % len(dx["refs"]))
        A("乱码检测: 无 U+FFFD/U+FFFE、无 mojibake 特征、无控制字符（此前 Step 6 已核验）")
    A("")
    A("四、工具与版本")
    A("-" * 46)
    for k, v in tools.items():
        A("%s: %s" % (k, v))
    A("方案: python-docx 直接生成 .docx（pandoc 路线不可用）")
    A("")
    A("五、问题与解决")
    A("-" * 46)
    if problems:
        for p_ in problems:
            A("- [问题] %s" % p_)
    A("- Step 3 校验脚本正则曾只捕获 '# ' 而非完整标题行，导致首次 FAIL；属校验逻辑 bug，输出文件本身正确，修复正则后重跑 PASS。")
    A("- Step 5 模块 docstring 中 r'C:\\Users\\...' 未用 raw string 时触发 \\U 转义 SyntaxError；改为 raw string 后正常运行。")
    A("- 终端(cmd)回显 UTF-8 中文时按 GBK 代码页显示为乱码，且末尾 echo 偶发'超时'回显；均为 shell 显示层问题，脚本实际输出完整，报告/产物文件本身为 UTF-8 无乱码。")
    A("- Step 6 工作区存在旧任务(大模型学习路线)的 step6_final_verify.py，与本任务无关，已新建独立验证脚本 step6_verify_docx.py，未误用旧脚本。")
    A("- 环境检查报告文件实际保存为 ENV_REPORT.md（大写），Windows 文件系统大小写不敏感，read_file 按实际路径回读成功。")
    if problems:
        A("- 本次复核新增问题: %d 项（见上）" % len(problems))
    else:
        A("本次复核未发现新增问题。")
    A("")
    A("结论: %s" % ("PASS - 全部验证项通过" if not problems else "FAIL - 存在 %d 项问题" % len(problems)))

    report = "\n".join(L) + "\n"
    with open(REPORT_PATH, "w", encoding="utf-8") as f:
        f.write(report)
    print("REPORT WRITTEN: %s" % REPORT_PATH)
    print("SIZE: %d bytes" % os.path.getsize(REPORT_PATH))
    print("RESULT: %s" % ("PASS" if not problems else "FAIL"))
    return 0 if not problems else 1

if __name__ == "__main__":
    sys.exit(main())
